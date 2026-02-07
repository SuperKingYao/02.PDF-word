#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转Word脚本 - 替代版本
使用 fitz (PyMuPDF) 和 python-docx 库将PDF文件转换为Word文档
"""

import os
import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: 缺少必要的库")
    print("请运行: pip install pymupdf python-docx")
    sys.exit(1)


def pdf_to_word_simple(pdf_path, output_path=None):
    """
    将PDF文件转换为Word文档（简化版 - 提取文本）
    
    参数:
        pdf_path (str): PDF文件路径
        output_path (str): 输出Word文件路径，默认为同名.docx文件
    
    返回:
        bool: 转换成功返回True，失败返回False
    """
    
    # 检查PDF文件是否存在
    if not os.path.exists(pdf_path):
        print(f"错误: PDF文件不存在 - {pdf_path}")
        return False
    
    # 如果没有指定输出路径，则使用同名.docx文件
    if output_path is None:
        pdf_name = Path(pdf_path).stem
        output_path = os.path.join(os.path.dirname(pdf_path), f"{pdf_name}.docx")
    
    try:
        print(f"开始转换: {pdf_path}")
        print(f"输出文件: {output_path}")
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        doc = Document()
        
        # 遍历每一页
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            
            if text.strip():
                # 添加页码标题
                heading = doc.add_heading(f'第 {page_num + 1} 页', level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加文本内容
                doc.add_paragraph(text)
                doc.add_paragraph()  # 添加空行
        
        # 保存Word文档
        doc.save(output_path)
        
        # 检查输出文件
        if os.path.exists(output_path):
            size = os.path.getsize(output_path) / 1024  # 转换为KB
            page_count = len(pdf_document)
            print(f"✓ 转换成功!")
            print(f"  页数: {page_count}")
            print(f"  文件大小: {size:.2f} KB")
            return True
        else:
            print("✗ 转换失败: 输出文件未生成")
            return False
            
    except Exception as e:
        print(f"✗ 转换过程中出错: {str(e)}")
        return False
    finally:
        try:
            pdf_document.close()
        except:
            pass


def pdf_to_word_with_images(pdf_path, output_path=None):
    """
    将PDF文件转换为Word文档（含图片版本）
    
    参数:
        pdf_path (str): PDF文件路径
        output_path (str): 输出Word文件路径，默认为同名.docx文件
    
    返回:
        bool: 转换成功返回True，失败返回False
    """
    
    # 检查PDF文件是否存在
    if not os.path.exists(pdf_path):
        print(f"错误: PDF文件不存在 - {pdf_path}")
        return False
    
    # 如果没有指定输出路径，则使用同名.docx文件
    if output_path is None:
        pdf_name = Path(pdf_path).stem
        output_path = os.path.join(os.path.dirname(pdf_path), f"{pdf_name}.docx")
    
    temp_dir = os.path.join(os.path.dirname(output_path), ".pdf_temp")
    
    try:
        print(f"开始转换: {pdf_path}")
        print(f"输出文件: {output_path}")
        
        # 创建临时目录用于存储页面图片
        os.makedirs(temp_dir, exist_ok=True)
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        doc = Document()
        
        # 遍历每一页
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # 添加页码标题
            heading = doc.add_heading(f'第 {page_num + 1} 页', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 将PDF页面渲染为图片
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))  # 1.5倍缩放
            image_path = os.path.join(temp_dir, f"page_{page_num}.png")
            pix.save(image_path)
            
            # 添加图片到Word
            doc.add_picture(image_path, width=Inches(6.5))
            
            # 提取文本
            text = page.get_text()
            if text.strip():
                doc.add_paragraph("【文本内容】")
                doc.add_paragraph(text)
            
            doc.add_paragraph()  # 添加空行
        
        # 保存Word文档
        doc.save(output_path)
        
        # 清理临时文件
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        # 检查输出文件
        if os.path.exists(output_path):
            size = os.path.getsize(output_path) / 1024  # 转换为KB
            page_count = len(pdf_document)
            print(f"✓ 转换成功!")
            print(f"  页数: {page_count}")
            print(f"  文件大小: {size:.2f} KB")
            return True
        else:
            print("✗ 转换失败: 输出文件未生成")
            return False
            
    except Exception as e:
        print(f"✗ 转换过程中出错: {str(e)}")
        return False
    finally:
        try:
            pdf_document.close()
        except:
            pass


def batch_convert(pdf_folder, output_folder=None, with_images=False):
    """
    批量转换PDF文件夹中的所有PDF文件
    
    参数:
        pdf_folder (str): 包含PDF文件的文件夹路径
        output_folder (str): 输出文件夹路径，默认为输入文件夹
        with_images (bool): 是否包含PDF页面的图片
    """
    
    if not os.path.isdir(pdf_folder):
        print(f"错误: 文件夹不存在 - {pdf_folder}")
        return
    
    if output_folder is None:
        output_folder = pdf_folder
    else:
        os.makedirs(output_folder, exist_ok=True)
    
    # 查找所有PDF文件
    pdf_files = list(Path(pdf_folder).glob("*.pdf"))
    
    if not pdf_files:
        print(f"未找到PDF文件 - {pdf_folder}")
        return
    
    print(f"找到 {len(pdf_files)} 个PDF文件，开始批量转换...\n")
    
    success_count = 0
    fail_count = 0
    
    convert_func = pdf_to_word_with_images if with_images else pdf_to_word_simple
    
    for pdf_file in pdf_files:
        output_path = os.path.join(output_folder, f"{pdf_file.stem}.docx")
        if convert_func(str(pdf_file), output_path):
            success_count += 1
        else:
            fail_count += 1
        print()
    
    print(f"批量转换完成! 成功: {success_count}, 失败: {fail_count}")


if __name__ == "__main__":
    
    # 使用示例
    if len(sys.argv) > 1:
        # 命令行模式
        pdf_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
        pdf_to_word_simple(pdf_path, output_path)
    else:
        # 交互模式
        print("=" * 50)
        print("PDF转Word转换工具 (替代版)")
        print("=" * 50)
        print("\n请选择转换模式:")
        print("1. 单个文件转换 (仅文本)")
        print("2. 单个文件转换 (含图片)")
        print("3. 批量转换文件夹 (仅文本)")
        print("4. 批量转换文件夹 (含图片)")
        
        choice = input("\n请输入选择 (1-4): ").strip()
        
        if choice == "1":
            pdf_path = input("请输入PDF文件路径: ").strip()
            output_path = input("请输入输出Word文件路径 (回车使用默认名称): ").strip()
            output_path = output_path if output_path else None
            pdf_to_word_simple(pdf_path, output_path)
            
        elif choice == "2":
            pdf_path = input("请输入PDF文件路径: ").strip()
            output_path = input("请输入输出Word文件路径 (回车使用默认名称): ").strip()
            output_path = output_path if output_path else None
            pdf_to_word_with_images(pdf_path, output_path)
            
        elif choice == "3":
            folder_path = input("请输入包含PDF文件的文件夹路径: ").strip()
            output_folder = input("请输入输出文件夹路径 (回车使用输入文件夹): ").strip()
            output_folder = output_folder if output_folder else None
            batch_convert(folder_path, output_folder, with_images=False)
            
        elif choice == "4":
            folder_path = input("请输入包含PDF文件的文件夹路径: ").strip()
            output_folder = input("请输入输出文件夹路径 (回车使用输入文件夹): ").strip()
            output_folder = output_folder if output_folder else None
            batch_convert(folder_path, output_folder, with_images=True)
            
        else:
            print("无效的选择!")
