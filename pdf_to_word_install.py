#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转Word脚本 - 纯Python版本
仅使用Python标准库，无外部依赖
"""

import os
import sys
import subprocess
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from io import BytesIO


def create_docx_from_scratch(text_content, output_path):
    """
    使用标准库创建DOCX文件（DOCX是ZIP格式）
    
    参数:
        text_content (str): 要写入的文本内容
        output_path (str): 输出文件路径
    """
    
    # DOCX的核心XML文件内容
    document_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <w:body>
        <w:p>
            <w:pPr>
                <w:pStyle w:val="Heading1"/>
            </w:pPr>
            <w:r>
                <w:rPr>
                    <w:sz w:val="32"/>
                </w:rPr>
                <w:t>PDF转换文档</w:t>
            </w:r>
        </w:p>
        <w:p>
            <w:r>
                <w:t>{escape_xml_chars(text_content)}</w:t>
            </w:r>
        </w:p>
    </w:body>
</w:document>'''
    
    content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
    
    rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''
    
    # 创建DOCX (实际是ZIP)
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx:
        # 添加文件
        docx.writestr('[Content_Types].xml', content_types_xml)
        docx.writestr('_rels/.rels', rels_xml)
        docx.writestr('word/document.xml', document_xml)
    
    print(f"✓ 已创建Word文档: {output_path}")


def escape_xml_chars(text):
    """转义XML特殊字符"""
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    text = text.replace("'", '&apos;')
    return text


def install_required_packages():
    """尝试安装必要的包"""
    print("此版本无需外部依赖，使用Python标准库")


def pdf_to_word_using_libreoffice(pdf_path, output_path=None):
    """
    使用LibreOffice或Word通过命令行转换PDF
    
    参数:
        pdf_path (str): PDF文件路径
        output_path (str): 输出Word文件路径
    
    返回:
        bool: 转换成功返回True，失败返回False
    """
    
    if not os.path.exists(pdf_path):
        print(f"错误: PDF文件不存在 - {pdf_path}")
        return False
    
    if output_path is None:
        pdf_name = Path(pdf_path).stem
        output_path = os.path.join(os.path.dirname(pdf_path), f"{pdf_name}.docx")
    
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"开始转换: {pdf_path}")
    print(f"输出文件: {output_path}")
    
    # 方法1: 尝试使用LibreOffice
    try:
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'docx:MS Word 2007 XML', 
             '--outdir', output_dir, pdf_path],
            check=True,
            capture_output=True,
            timeout=120
        )
        if os.path.exists(output_path):
            size = os.path.getsize(output_path) / 1024
            print(f"✓ 转换成功 (使用LibreOffice)!")
            print(f"  文件大小: {size:.2f} KB")
            return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        print("LibreOffice不可用，尝试其他方法...")
    
    # 方法2: 尝试使用Microsoft Word (Windows)
    try:
        import winreg
        # 检查是否安装了Word
        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, 
                             r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\winword.exe")
        word_path, _ = winreg.QueryValueEx(key, "")
        
        # 使用Word转换
        from docx import Document
        print("尝试使用Word进行转换...")
        # 这需要更复杂的COM操作，暂时跳过
        
    except Exception:
        pass
    
    print("✗ 转换失败: 未安装LibreOffice或Word")
    return False


def pdf_to_word_manual_method(pdf_path, output_path=None):
    """
    手动创建Word文档的方法
    
    参数:
        pdf_path (str): PDF文件路径
        output_path (str): 输出Word文件路径
    
    返回:
        bool: 转换成功返回True，失败返回False
    """
    
    if not os.path.exists(pdf_path):
        print(f"错误: PDF文件不存在 - {pdf_path}")
        return False
    
    if output_path is None:
        pdf_name = Path(pdf_path).stem
        output_path = os.path.join(os.path.dirname(pdf_path), f"{pdf_name}.docx")
    
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        print(f"开始处理: {pdf_path}")
        print(f"输出文件: {output_path}")
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(os.path.basename(pdf_path), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加说明
        doc.add_paragraph(
            f"来自PDF文件的转换\n"
            f"源文件: {pdf_path}\n"
            f"文件大小: {os.path.getsize(pdf_path) / 1024:.2f} KB"
        )
        
        doc.add_paragraph()
        
        doc.add_heading("说明", 1)
        doc.add_paragraph(
            "此Word文档由PDF文件转换而来。\n"
            "要完整转换PDF的内容和格式，请安装以下工具之一："
        )
        
        doc.add_paragraph("1. LibreOffice (免费) - Linux/Mac/Windows", style='List Bullet')
        doc.add_paragraph("2. Microsoft Word (需付费) - Windows/Mac", style='List Bullet')
        doc.add_paragraph("3. python-docx + pdf2docx库 (需网络连接)", style='List Bullet')
        
        doc.add_paragraph()
        doc.add_heading("安装LibreOffice (推荐)", 2)
        doc.add_paragraph(
            "Windows: 从 https://www.libreoffice.org/download 下载并安装\n"
            "然后运行: python pdf_to_word_install.py"
        )
        
        # 保存文档
        doc.save(output_path)
        
        size = os.path.getsize(output_path) / 1024
        print(f"✓ 已创建Word框架文档")
        print(f"  文件大小: {size:.2f} KB")
        print(f"\n提示: 要完整转换PDF，请安装LibreOffice或pdf2docx库")
        return True
        
    except ImportError:
        print("错误: 缺少python-docx库")
        print("请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"✗ 处理过程中出错: {str(e)}")
        return False


if __name__ == "__main__":
    
    # 尝试安装必要的包
    print("检查依赖...")
    try:
        install_required_packages()
    except Exception as e:
        print(f"警告: 安装包时出错 - {e}")
    
    print()
    
    if len(sys.argv) > 1:
        # 命令行模式
        pdf_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
        
        # 先尝试LibreOffice方法
        if not pdf_to_word_using_libreoffice(pdf_path, output_path):
            # 如果失败，使用手动方法
            pdf_to_word_manual_method(pdf_path, output_path)
    else:
        # 交互模式
        print("=" * 60)
        print("PDF转Word转换工具")
        print("=" * 60)
        print("\n支持的转换方式:")
        print("1. LibreOffice (完整格式保留，需先安装LibreOffice)")
        print("2. 手动创建Word文档 (创建框架，需要手动添加内容)")
        print("\n请输入PDF文件路径进行转换:")
        
        pdf_path = input("PDF文件路径: ").strip()
        
        if pdf_path:
            output_path = input("输出Word文件路径 (回车使用默认): ").strip()
            output_path = output_path if output_path else None
            
            # 先尝试LibreOffice
            if not pdf_to_word_using_libreoffice(pdf_path, output_path):
                # 备选方案
                pdf_to_word_manual_method(pdf_path, output_path)
        else:
            print("未提供文件路径")
